"""
DOCX parser — extracts a structured section map from a .docx byte stream.

Returns a dict with three top-level keys::

    __meta__     document fingerprint and summary statistics
    sections     dict of normalised heading → section envelope
    raw_text     full document text for downstream AI extraction

Pipeline stages:
  1. Document fingerprinting  — page_title, university_name, program_name
  2. Heading normalisation    — strip uni/program prefix, junk words
  3. Content type detection   — faq / bullets / table / mixed / text
  4. Table subtype detection  — semester / horizontal / uniform / standard
  5. Final output assembly    — wrapped in standard envelope per section

All detection is deterministic — no API calls.
"""

from __future__ import annotations

import io
import re
from collections import Counter
from typing import Any

from docx import Document


# ━━━━━━━━━━━━━━━━━━━━━━━━  COMPILED PATTERNS  ━━━━━━━━━━━━━━━━━━━━━━━━

# Program keyword boundary — splits university name from the rest of a title
_PROGRAM_KW_RE = re.compile(
    r"\b(?:online|distance|mba|bba|bca|bsc|b\.?\s*sc|mca|"
    r"m\.?\s*com|mcom|b\.?\s*com|bcom|b\.?\s*tech|m\.?\s*tech|"
    r"ma\s|phd|pgdm)\b",
    re.IGNORECASE,
)

# Common section-like suffixes at the end of a document title
_SECTION_SUFFIX_RE = re.compile(
    r"\b(?:"
    r"course\s+details?|fee\s+structures?|fees?\s+details?|"
    r"admission\s+process|eligibility(?:\s+criteria)?|"
    r"syllabus|curriculum|career\s+(?:opportunities|scope)|"
    r"job\s+(?:opportunities|profiles?)|placements?|"
    r"highlights?|overview|about(?:\s+(?:course|program))?|faqs?|"
    r"specializations?|advantages|benefits|"
    r"key\s+features|emi(?:\s+(?:details?|options?))?|"
    r"salary|scope|details?|structure|process|criteria|"
    r"comparisons?|rankings?|reviews?|duration"
    r")\s*$",
    re.IGNORECASE,
)

# Leading junk words stripped after removing university/program prefix
_LEADING_JUNK_RE = re.compile(
    r"^[\s,\-–—:]*\b(?:online|in|for|of|the|and|&)\b[\s,\-–—:]*",
    re.IGNORECASE,
)

# Bullet markers at the start of a line
_BULLET_RE = re.compile(
    r"^(?:[•●○▪▸▹►–—\-]\s*|\d+[.)]\s+|[a-zA-Z][.)]\s+|"
    r"[\u2022\u2023\u25E6\u2043\u2219]\s*)"
)

# Year / Semester patterns for semester-table detection
_YEAR_RE = re.compile(r"\bYear\s+[IVXivx\d]+\b", re.IGNORECASE)
_SEM_RE = re.compile(r"\b(?:Semester|Sem\.?)\s*[IVXivx\d]+\b", re.IGNORECASE)

# Horizontal-table detection: label keywords in header row
_LABEL_KW_RE = re.compile(
    r"(?:months?|emi|plan|option|semester|year|type|duration|fee|payment)",
    re.IGNORECASE,
)
# Horizontal-table detection: value indicators in data row
_VALUE_IND_RE = re.compile(r"(?:INR|₹|Rs\.?|\d{2,}|%|/\-)", re.IGNORECASE)

# Header normalisation helpers
_NON_ALNUM_RE = re.compile(r"[^\w\s]")
_MULTI_USCORE_RE = re.compile(r"_+")


# ────────────────────────  SECTION KEYWORDS ────────────────────────

# Terminal section keywords — these appear at the RIGHT end of institution headings.
# Ordered longest-first so the most specific keyword wins.
# Key insight: headings like "XYZ University Online MBA [keyword]" always end
# with the section name. rfind() on this list finds the correct anchor.
SECTION_KEYWORDS: list[str] = [
    # ── Multi-word (longest first) ──
    "specialization wise fees", "specialization-wise fees",
    "job roles & salary", "job roles and salary",
    "frequently asked questions",
    "approvals and accreditations",
    "sample certificate photo", "sample certificate",
    "career and placements", "career & placements",
    "placement partners", "top recruiters", "hiring partners",
    "faculty members", "meet the faculty", "our faculty",
    "student reviews", "what students say", "alumni speak",
    "program highlights", "course highlights", "key highlights",
    "program overview", "program details",
    "about the university", "about the course", "about the program",
    "about the specialization",
    "course curriculum", "course overview", "course details",
    "syllabus/curriculum", "syllabus / curriculum",
    "fee structure", "fee & emi", "fees & emi",
    "emi details", "emi options", "fee and emi",
    "fee payment", "fee details", "course fee",
    "admission process", "admission procedure", "how to apply",
    "enrollment process",
    "eligibility criteria", "who can apply",
    "entry requirements", "admission requirements",
    "examination pattern", "examination process",
    "exam pattern", "exam process",
    "assessment pattern",
    "available specializations", "other specializations",
    "explore specializations", "related specializations",
    "degree certificate",
    "no cost emi", "no-cost emi", "easy emi",
    "short description", "exam schedule",
    "job profiles", "job roles", "career opportunities",
    "placement support", "career support",
    "career outcomes",
    "key features", "why this program",
    "programs offered", "courses offered",
    "why choose", "why opt for", "key benefits",
    "quick facts", "key facts",
    "course facts",
    "our recruiters", "recruiting companies",
    "teaching faculty",
    # ── Single-word (shorter; must come after all multi-word) ──
    "accreditations", "accreditation",
    "approvals", "recognition", "certifications",
    "highlights", "overview", "syllabus", "curriculum",
    "admission", "eligibility",
    "placements", "placement", "faculty", "reviews",
    "testimonials", "faqs", "faq",
    "specializations", "certificate", "emi",
    "introduction", "description", "examination",
    "courses", "programs",
    "details", "info", "information",
    "pros", "benefits", "advantages",
    "facts", "about",
]

# Noise token patterns — these appear on the LEFT in institution headings.
# Covers: degree-level words AND institution-type words.
# Applied LEFT-TO-RIGHT in one pass using a single combined pattern.
# Does NOT anchor to ^ because an unknown proper noun may precede the noise.
_NOISE_TOKEN_RE = re.compile(
    r"\b(?:"
    # Institution type words
    r"university|college|institute|institution|academy|school|polytechnic"
    r"|department|dept"
    r"|"
    # Delivery mode words
    r"online|distance|hybrid|blended"
    r"|"
    # Degree / program type words
    r"mba|mca|bba|bca|b\.?\s*tech|m\.?\s*tech|bsc|b\.?\s*sc"
    r"|msc|m\.?\s*sc|b\.?\s*com|bcom|m\.?\s*com|mcom"
    r"|ba|ma|llb|llm|phd|pgdm|diploma|pg"
    r"|"
    # Structural connector words (only as prefix noise)
    r"program|course|page"
    r")\b",
    re.IGNORECASE,
)

# Stray leading conjunctions / prepositions left after noise removal
_JUNK_PREFIX_RE = re.compile(r"^[\s\-–—,:.]+\b(in|of|for|and|&|the|a|an)\b[\s\-–—,:.]*", re.IGNORECASE)


def strip_university_prefix(heading: str) -> str:
    """Remove any institution/program name prefix from a document heading.

    Works for ANY institution — zero hardcoded names.

    Strategy
    --------
    1. **rfind keyword anchor** (primary):
       Scan SECTION_KEYWORDS using *rfind* (rightmost match).  The keyword
       is always at the RIGHT end of institutional headings like
       ``"XYZ University Online MBA [keyword]"``.  Returning from the
       rightmost keyword position gives the cleanest result.

    2. **Noise-token scrubbing** (secondary):
       Remove institution-type and degree-type noise tokens from the
       cleaned text, then strip stray leading conjunctions.

    3. **Short result fast-return**:
       If cleaned text ≤ 5 words, return it directly.

    4. **Last-40%-words fallback**:
       For long headings where no keyword was found, return the last 40 %
       of words.

    Examples::

        "Mody University Online details"         → "details"
        "about Mody University Online"           → "about"
        "Some University course facts"           → "course facts"
        "Mody University Online pros"            → "pros"
        "XYZ College Online BCA Syllabus"        → "Syllabus"
        "abc university admission process"       → "admission process"
        "Greenfield Academy Diploma syllabus"    → "syllabus"
    """
    original = heading.strip()
    lower = original.lower()

    # ── Phase 1: rfind keyword anchor ───────────────────────────────────────
    # Collect ALL keyword matches: (start_idx, end_idx, keyword)
    # Then select the one whose start is furthest right.
    # Tie-break: prefer longer keyword (to pick "syllabus/curriculum" over "curriculum").
    # Overlap rule: if keyword B starts inside keyword A's span, discard B — A is longer.
    all_matches: list[tuple[int, int, str]] = []
    for keyword in SECTION_KEYWORDS:
        idx = lower.rfind(keyword)
        if idx == -1:
            continue
        all_matches.append((idx, idx + len(keyword), keyword))

    # Remove dominated matches: if match B is completely contained in match A, drop B
    # (sort by start desc, then end desc to process rightmost/longest first)
    all_matches.sort(key=lambda m: (m[0], -(m[1] - m[0])))  # start asc, then longest first

    # Among matches that don't overlap each other, find the rightmost-starting one
    best: tuple[int, int, str] | None = None
    for m in all_matches:
        if best is None:
            best = m
            continue
        m_start, m_end, m_kw = m
        b_start, b_end, b_kw = best
        # If m starts after best ends → m is further right → take m
        if m_start >= b_end:
            best = m
        # If m overlaps best (m starts inside best's span) → m is a sub-keyword of best → skip
        # If m starts at same position as best → prefer longer (already sorted that way)

    if best is not None:
        best_idx, best_end, best_kw = best
        result = original[best_idx:].strip()

        # Guard: keyword was at position 0 AND heading is multi-word
        # e.g. "about mody university online" → keyword "about" found at 0
        # → scrub noise tokens from the full string then return
        if best_idx == 0 and len(original.split()) > 1:
            scrubbed = _NOISE_TOKEN_RE.sub(" ", lower).strip()
            scrubbed = _JUNK_PREFIX_RE.sub("", scrubbed).strip()
            scrubbed = re.sub(r"\s{2,}", " ", scrubbed).strip()
            # Only use scrubbed if it's shorter (noise was actually removed)
            # and still contains the keyword
            if (scrubbed and len(scrubbed) < len(lower)
                    and best_kw in scrubbed):
                # If scrubbed is just the keyword alone, great; otherwise take it
                return scrubbed.strip()
            # If scrubbed removed the keyword too (e.g. "course" got removed),
            # fall back to the keyword itself
            if scrubbed and best_kw not in scrubbed:
                return best_kw.strip()
        return result

    # ── Phase 2: noise-token scrubbing ──────────────────────────────────────
    scrubbed = _NOISE_TOKEN_RE.sub(" ", lower).strip()
    scrubbed = _JUNK_PREFIX_RE.sub("", scrubbed).strip()
    scrubbed = re.sub(r"\s{2,}", " ", scrubbed).strip()

    if scrubbed and len(scrubbed.split()) <= 5:
        return scrubbed.strip()

    # ── Phase 3: last 40 % of words fallback ────────────────────────────────
    words = original.split()
    if len(words) > 3:
        cutoff = max(1, int(len(words) * 0.4))
        return " ".join(words[-cutoff:]).strip()

    return original.strip()





# ━━━━━━━━━━━━━━━━━━━━━━━━━━  PUBLIC API  ━━━━━━━━━━━━━━━━━━━━━━━━━━


def parse_docx(file_bytes: bytes) -> dict[str, dict[str, Any]]:
    """Return a structured section map extracted from the given *.docx* bytes.

    The output has three top-level keys: ``__meta__``, ``sections``,
    and ``raw_text``.  See module docstring for details.
    """

    doc = Document(io.BytesIO(file_bytes))

    # Pre-index elements for O(1) lookup while walking the body
    para_map: dict[int, Any] = {id(p._element): p for p in doc.paragraphs}
    table_map: dict[int, Any] = {id(t._element): t for t in doc.tables}

    # ── STAGE 1: Document Fingerprinting ──────────────────────
    meta = _fingerprint(doc, para_map)

    # ── Walk the body in document order → raw sections ────────
    raw_sections = _walk_body(doc, para_map, table_map)

    # ── Collect raw_text ──────────────────────────────────────
    raw_text = _extract_raw_text(doc)

    # ── STAGES 2–5: normalise, detect, build envelopes ────────
    sections: dict[str, dict[str, Any]] = {}
    detected_headings: list[str] = []
    table_count = 0
    has_faq = False

    for heading, content_items in raw_sections.items():
        detected_headings.append(heading)

        # Stage 2: heading normalisation
        stripped = _normalize_heading(
            heading, meta["university_name"], meta["program_name"]
        )

        # De-duplicate keys (append _2, _3, … if collision)
        if stripped in sections:
            n = 2
            while f"{stripped}_{n}" in sections:
                n += 1
            stripped = f"{stripped}_{n}"

        # Count tables in this section
        section_tables = [i for i in content_items if i["type"] == "table"]
        table_count += len(section_tables)

        # Stage 3 + 4: detect content type, classify tables, parse
        content_type, parsed_content, table_subtype = _detect_and_parse(
            content_items
        )
        if content_type == "faq":
            has_faq = True

        # Stage 5: wrap in standard envelope
        envelope = _build_envelope(
            stripped, heading, content_type, table_subtype, parsed_content
        )
        sections[stripped] = envelope

    # Finalise meta
    meta["section_count"] = len(sections)
    meta["table_count"] = table_count
    meta["has_faq"] = has_faq
    meta["detected_headings"] = detected_headings

    return {
        "__meta__": meta,
        "sections": sections,
        "raw_text": raw_text,
    }


# ━━━━━━━━━━━━━━  STAGE 1 — DOCUMENT FINGERPRINTING  ━━━━━━━━━━━━━━


def _fingerprint(doc, para_map: dict[int, Any]) -> dict[str, Any]:
    """Extract *page_title*, *university_name*, and *program_name*.

    page_title = the first non-empty paragraph that appears BEFORE
                 the first heading and is not a heading style itself.
                 Falls back to the first heading text if none found.
    """
    page_title = ""
    first_heading = ""
    hit_first_heading = False

    for child in doc.element.body:
        if not child.tag.endswith("}p"):
            continue
        para = para_map.get(id(child))
        if para is None:
            continue
        text = para.text.strip()
        if not text:
            continue

        if _is_heading(para):
            if not first_heading:
                first_heading = text
            hit_first_heading = True
            # Stop scanning — page_title must appear before first heading
            break

        # Candidate paragraph before the first heading.
        # Accept it as page_title (it is not a heading style).
        if not page_title:
            page_title = text
            # Don't break yet — keep scanning until we confirm there
            # IS a heading after it (or until document ends).

    if not page_title:
        page_title = first_heading or ""

    uni = _extract_university_name(page_title)
    prog = _extract_program_name(page_title, uni)

    return {
        "page_title": page_title,
        "university_name": uni,
        "program_name": prog,
    }


def _extract_university_name(title: str) -> str:
    """Everything *before* the first program keyword in the title.

    Example::

        "Chandigarh University Online MBA in Finance"
        → "Chandigarh University"
    """
    m = _PROGRAM_KW_RE.search(title)
    if m:
        return title[: m.start()].strip()
    return ""


def _extract_program_name(title: str, uni_name: str) -> str:
    """Everything between *university_name* and the trailing section suffix.

    Example::

        title    = "Chandigarh University Online MBA in Finance Course Details"
        uni_name = "Chandigarh University"
        → "Online MBA in Finance"
    """
    remainder = title
    if uni_name:
        pat = re.compile(re.escape(uni_name), re.IGNORECASE)
        remainder = pat.sub("", remainder, count=1).strip()

    if not remainder:
        return title

    # Strip section suffix from the end
    m = _SECTION_SUFFIX_RE.search(remainder)
    if m:
        prog = remainder[: m.start()].strip()
        prog = re.sub(r"[\s,\-–—:]+$", "", prog).strip()
        if prog:
            return prog

    return remainder


# ━━━━━━━━━━━━━━  STAGE 2 — HEADING NORMALISATION  ━━━━━━━━━━━━━━


def _normalize_heading(heading: str, uni_name: str, prog_name: str) -> str:
    """Normalise a document heading to a clean ``stripped_key``.

    Steps:
      A. Strip university name prefix
      B. Strip program name
      C. Strip leading/trailing junk (online, in, for, of, -, ,)
      D. Lowercase and trim
      E. If result < 3 chars, fall back to strip_university_prefix()
         then finally the original heading lowercased
    """
    result = heading

    # A — strip university name
    if uni_name:
        pat = re.compile(re.escape(uni_name), re.IGNORECASE)
        result = pat.sub("", result, count=1).strip()

    # B — strip program name
    if prog_name:
        pat = re.compile(re.escape(prog_name), re.IGNORECASE)
        result = pat.sub("", result, count=1).strip()

    # C — strip leading junk words (repeat until stable)
    prev = None
    while result != prev:
        prev = result
        result = _LEADING_JUNK_RE.sub("", result).strip()
    # Also strip trailing junk characters
    result = re.sub(r"[\s,\-–—:]+$", "", result).strip()

    # D — lowercase
    result = result.lower().strip()

    # E — safety: never return empty or too-short key
    if len(result) < 3:
        # Try the generic keyword-based stripper before giving up
        fallback = strip_university_prefix(heading).lower().strip()
        if len(fallback) >= 3:
            return fallback
        result = heading.lower().strip()

    return result


# ━━━━━━━━━━━━━━  DOCUMENT BODY WALK + HEADING DETECTION  ━━━━━━━━━━━━━━


# Keywords that must appear for a text to be a genuine section heading.
# List items / company names / proper nouns usually don't contain these.
HEADING_MUST_HAVE_KEYWORDS: list[str] = [
    "about", "overview", "introduction", "accreditation", "eligibility",
    "admission", "process", "fee", "emi", "payment", "syllabus", "curriculum",
    "placement", "partner", "recruiter", "job", "role", "salary", "career",
    "review", "testimonial", "faq", "question", "fact", "highlight",
    "specialization", "program", "course", "certificate", "exam",
    "university", "college", "institute", "why", "choose", "apply",
    "structure", "pattern", "criteria", "requirement", "details",
]


def is_real_section_heading(text: str, style_name: str = "") -> bool:
    """Return True only if this text is a genuine section heading.

    Distinguishes section headings from list items and content text
    that happens to be bold (e.g. company names inside a placement
    section, specialization names inside a bullet list).

    Rules
    -----
    * Explicit heading styles (Heading 1/2/3) always qualify.
    * Single words → always reject (never a section heading).
    * 2–4 words → must contain at least one heading keyword.
    * 5+ words with bold formatting → likely a heading (pass through).
    """
    text = text.strip()
    if not text:
        return False

    # Explicit heading styles always qualify regardless of content
    if style_name and any(
        h in style_name.lower() for h in ("heading", "h1", "h2", "h3", "title")
    ):
        return True

    lower = text.lower()
    word_count = len(text.split())

    # Single words are almost never section headings
    if word_count == 1:
        return False

    # 2–4 words: require at least one recognisable heading keyword
    if word_count <= 4:
        return any(kw in lower for kw in HEADING_MUST_HAVE_KEYWORDS)

    # 5+ words: pass through (handled by bold/style detection upstream)
    return True


def _walk_body(
    doc,
    para_map: dict[int, Any],
    table_map: dict[int, Any],
) -> dict[str, list[dict[str, Any]]]:
    """Walk the document body in order and split into heading → content items.

    Each content item is either::

        {"type": "paragraph", "text": str, "is_list_style": bool}
        {"type": "table",     "rows": list[list[str]]}

    Sections with no content items are discarded.
    Orphaned sections (single-word headings, company names, etc.) have
    their content merged back into the preceding real section.
    """
    sections: dict[str, list[dict[str, Any]]] = {}
    current_heading = "Introduction"
    current_content: list[dict[str, Any]] = []

    for child in doc.element.body:
        eid = id(child)

        # ── paragraph ──
        if child.tag.endswith("}p"):
            para = para_map.get(eid)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue

            style_name = ""
            if para.style and para.style.name:
                style_name = para.style.name

            if _is_heading(para):
                if is_real_section_heading(text, style_name):
                    # Flush previous section only if it has content
                    if current_content:
                        sections[current_heading] = current_content
                    current_heading = text
                    current_content = []
                else:
                    # Not a real section heading — treat as content
                    # and absorb into the current section as a list item.
                    # (e.g. "Amazon", "Human Resource Management" inside a list)
                    current_content.append(
                        {
                            "type": "paragraph",
                            "text": text,
                            "is_list_style": True,  # treat as list item
                            "style_name": style_name,
                        }
                    )
            else:
                is_list = False
                sn = style_name.lower()
                is_list = any(
                    kw in sn for kw in ("list", "bullet", "number")
                )
                current_content.append(
                    {
                        "type": "paragraph",
                        "text": text,
                        "is_list_style": is_list,
                        "style_name": style_name,
                    }
                )

        # ── table ──
        elif child.tag.endswith("}tbl"):
            table = table_map.get(eid)
            if table is None:
                continue
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(c for c in cells):  # skip fully empty rows
                    rows.append(cells)
            if rows:
                current_content.append({"type": "table", "rows": rows})

    # Flush last section only if it has content
    if current_content:
        sections[current_heading] = current_content

    return sections


def _is_heading(para) -> bool:
    """Return True if *para* looks like a section heading.

    Detection priority:
      1. Paragraph style name contains "Heading" or "Title"
      2. All runs are bold AND the largest font-size ≥ 14 pt
      3. All runs are bold, text < 100 chars, no explicit font size
         (catches headings inheriting size from the Word style)
      4. ALL CAPS short line (≤ 8 words)
      5. Title Case short line with no trailing punctuation (≤ 10 words)
    """
    # 1. Style-based
    if para.style and para.style.name:
        name = para.style.name
        if "Heading" in name or "Title" in name:
            return True

    text = para.text.strip()
    if not text:
        return False

    # 2+3. Font-based heuristic
    runs = [r for r in para.runs if r.text.strip()]
    if runs and all(_run_is_bold(r) for r in runs):
        max_pt = 0.0
        for r in runs:
            size = r.font.size
            if size is not None:
                pt_val = size.pt
                if pt_val > max_pt:
                    max_pt = pt_val

        if max_pt >= 14:
            return True

        # Fallback: all bold + short text + no explicit font size
        if max_pt == 0.0 and len(text) < 100:
            return True

    # 4. ALL CAPS short line
    if text.isupper() and len(text.split()) <= 8 and len(text) >= 3:
        return True

    # 5. Title Case + short + no trailing sentence punctuation
    if (
        not text.endswith(('.', ',', ';', '?'))
        and text.istitle()
        and len(text.split()) <= 10
        and len(text) < 100
        and len(text) >= 5
    ):
        return True

    return False


def _run_is_bold(run) -> bool:
    """Check if a run is bold, accounting for inherited character styles."""
    if run.bold is True:
        return True
    if run.bold is None:
        try:
            style = run.style
            if style and style.font and style.font.bold:
                return True
        except Exception:
            pass
    return False


# ━━━━━━━━━━━━━━  STAGE 3 — CONTENT TYPE DETECTION  ━━━━━━━━━━━━━━


def _detect_and_parse(
    content_items: list[dict[str, Any]],
) -> tuple[str, Any, str | None]:
    """Detect content type, parse content, determine table subtype.

    Returns ``(content_type, parsed_content, table_subtype)``.

    *table_subtype* is non-None only when *content_type* is ``"table"``.
    """
    paragraphs = [i for i in content_items if i["type"] == "paragraph"]
    tables = [i for i in content_items if i["type"] == "table"]
    texts = [p["text"] for p in paragraphs]

    # ── Bullet-like detection ──
    bullet_by_regex = sum(1 for t in texts if _BULLET_RE.match(t))
    bullet_by_style = sum(1 for p in paragraphs if p.get("is_list_style"))
    total_bullets = max(bullet_by_regex, bullet_by_style)
    is_mostly_bullets = bool(texts and total_bullets / len(texts) >= 0.6)

    # ── FAQ detection: alternating ListNumber / ListBullet styles ──
    #    Question = ListNumber, Answer = ListBullet.
    #    Check BEFORE the generic bullet/question-mark FAQ path.
    if len(paragraphs) >= 4:
        faq_by_style = _detect_faq_by_alternating_styles(paragraphs)
        if faq_by_style:
            return ("faq", faq_by_style, None)

    # TYPE 1: FAQ — bullet points with at least 2 questions ("?" heuristic)
    if is_mostly_bullets and len(texts) >= 4:
        q_count = sum(
            1
            for i in range(0, len(texts), 2)
            if texts[i].rstrip().endswith("?")
        )
        if q_count >= 2:
            faq = _build_faq(texts)
            if faq:
                return ("faq", faq, None)

    # TYPE 1b: FAQ — plain alternating paragraphs (no list styles, no tables)
    #   Questions at indices 0, 2, 4, … end with "?"
    #   Answers at indices 1, 3, 5, … do not.
    if not tables and len(texts) >= 4:
        has_any_list = any(p.get("is_list_style") for p in paragraphs)
        if not has_any_list:
            even_indices = range(0, len(texts), 2)
            q_count = sum(
                1 for i in even_indices if texts[i].rstrip().endswith("?")
            )
            if q_count / len(list(even_indices)) > 0.5:
                faq = _build_faq(texts)
                if faq:
                    return ("faq", faq, None)

    # TYPE 2: Bullets — list items, no tables
    if is_mostly_bullets and not tables:
        cleaned = [_BULLET_RE.sub("", t).strip() for t in texts]
        cleaned = [c for c in cleaned if c]
        if not cleaned:
            cleaned = [t.strip() for t in texts if t.strip()]
        return ("bullets", cleaned, None)

    # TYPE 3: Table — exactly one table, no paragraphs
    if len(tables) == 1 and not paragraphs:
        subtype, data = _parse_table_with_subtype(tables[0]["rows"])
        return ("table", data, subtype)

    # TYPE 4: Mixed — paragraphs AND tables, or multiple tables no paragraphs
    if tables:
        intro_html = _texts_to_html(texts) if texts else None
        parsed_tables: list[dict[str, Any]] = []
        for t in tables:
            subtype, data = _parse_table_with_subtype(t["rows"])
            parsed_tables.append({"table_subtype": subtype, "data": data})
        mixed_content = {"intro_text": intro_html, "tables": parsed_tables}
        return ("mixed", mixed_content, None)

    # TYPE 5: Text — only paragraphs, no tables, no bullets
    if texts:
        return ("text", _texts_to_html(texts), None)

    return ("text", "", None)


def _detect_faq_by_alternating_styles(
    paragraphs: list[dict[str, Any]],
) -> list[dict[str, str]] | None:
    """Detect FAQ from alternating ListNumber (Q) / ListBullet (A) styles.

    Returns a ``[{question, answer}]`` list if the pattern is found,
    otherwise ``None``.
    """
    styles = [p.get("style_name", "").lower() for p in paragraphs]

    # Count how many pairs follow the ListNumber → ListBullet pattern
    pair_count = 0
    i = 0
    while i + 1 < len(styles):
        is_q = "number" in styles[i] and "list" in styles[i]
        is_a = "bullet" in styles[i + 1] and "list" in styles[i + 1]
        if is_q and is_a:
            pair_count += 1
            i += 2
        else:
            i += 1

    if pair_count < 2:
        return None

    # Build Q/A pairs
    faq: list[dict[str, str]] = []
    i = 0
    while i + 1 < len(paragraphs):
        sn_q = paragraphs[i].get("style_name", "").lower()
        sn_a = paragraphs[i + 1].get("style_name", "").lower()

        is_q = "number" in sn_q and "list" in sn_q
        is_a = "bullet" in sn_a and "list" in sn_a

        if is_q and is_a:
            faq.append({
                "question": paragraphs[i]["text"].strip(),
                "answer": paragraphs[i + 1]["text"].strip(),
            })
            i += 2
        else:
            # Absorb stray paragraph into previous answer
            if faq:
                faq[-1]["answer"] += " " + paragraphs[i]["text"].strip()
            i += 1

    return faq if faq else None


# ━━━━━━━━━━━━━━  STAGE 4 — TABLE SUBTYPE DETECTION & PARSING  ━━━━━━━━━━━━━━


def _parse_table_with_subtype(
    rows: list[list[str]],
) -> tuple[str, Any]:
    """Detect the table subtype and return ``(subtype, parsed_data)``.

    Subtypes tried in order: semester → horizontal → uniform → standard.
    If a subtype detector fires but parsing fails, falls through to
    the next candidate.  Never raises.
    """
    if not rows:
        return ("standard_table", [])

    try:
        # SUBTYPE A: semester_table
        if _is_semester_table(rows):
            data = _parse_semester_table(rows)
            if data:
                return ("semester_table", data)

        # SUBTYPE B: horizontal_table
        if _is_horizontal_table(rows):
            data = _parse_horizontal_table(rows)
            if data:
                return ("horizontal_table", data)

        # SUBTYPE C: uniform_table
        if _is_uniform_table(rows):
            data = _parse_uniform_table(rows)
            if data:
                return ("uniform_table", data)
    except Exception:
        pass  # Fall through to standard_table

    # SUBTYPE D: standard_table (default)
    return ("standard_table", _parse_standard_table(rows))


# ── subtype detection helpers ──


def _is_semester_table(rows: list[list[str]]) -> bool:
    """True if any row contains a Year or Semester marker.

    Checks:
      - First column contains "Year I", "Semester I", etc.
      - Both cells in a row have the same text starting with Year/Semester
        (merged-cell pattern)
    """
    for row in rows:
        for cell in row:
            if _YEAR_RE.search(cell) or _SEM_RE.search(cell):
                return True
    return False


def _is_horizontal_table(rows: list[list[str]]) -> bool:
    """True if table has exactly 2 data rows — labels then values.

    Before checking, skips rows that are merged-cell title rows
    (all cells contain the same text) or single-cell spanning rows.
    """
    # Filter out merged / title rows
    filtered = _skip_merged_title_rows(rows)

    if len(filtered) != 2:
        return False
    headers, values = filtered[0], filtered[1]
    label_hits = sum(1 for h in headers if _LABEL_KW_RE.search(h))
    value_hits = sum(1 for v in values if _VALUE_IND_RE.search(v))
    return label_hits >= 1 and value_hits >= 1


def _skip_merged_title_rows(
    rows: list[list[str]],
) -> list[list[str]]:
    """Return *rows* with merged / title rows removed.

    A row is considered a merged title row if:
      - All cells contain the exact same text, OR
      - Only one cell has content (single spanning cell)
    """
    out: list[list[str]] = []
    for row in rows:
        stripped = [c.strip() for c in row]
        non_empty = [c for c in stripped if c]

        # All cells identical → merged title row → skip
        if non_empty and len(set(non_empty)) == 1 and len(non_empty) == len(stripped):
            continue

        # Only one cell has content (spans full width) → skip
        if len(non_empty) == 1 and len(stripped) > 1:
            continue

        out.append(row)
    return out


def _is_uniform_table(rows: list[list[str]]) -> bool:
    """True if 2-column table where >70 % of column-2 values are identical."""
    if len(rows) < 3:
        return False
    if any(len(r) != 2 for r in rows):
        return False
    col2 = [r[1].strip() for r in rows[1:] if len(r) > 1 and r[1].strip()]
    if not col2:
        return False
    top_count = Counter(col2).most_common(1)[0][1]
    return top_count / len(col2) > 0.7


# ── subtype parsing helpers ──


def _parse_semester_table(rows: list[list[str]]) -> dict[str, Any]:
    """Parse a Year/Semester curriculum table into nested structure.

    Output::

        {
          "Year I": {
            "Semester I": ["subject1", "subject2"],
            "Semester II": ["subject3", "subject4"]
          },
          "Year II": { ... }
        }

    Partially-successful parses are returned as-is (partial > empty).
    """
    result: dict[str, dict[str, list[str]]] = {}
    current_year = ""
    current_sems: list[str] = ["", ""]  # [left_col, right_col]

    for row in rows:
        if not row:
            continue

        # Pad to at least 2 columns for uniform access
        padded = list(row)
        while len(padded) < 2:
            padded.append("")

        c0, c1 = padded[0].strip(), padded[1].strip()

        # ── Year header: both cells identical AND contains "Year" ──
        if c0 and c0 == c1 and _YEAR_RE.search(c0):
            current_year = c0
            result.setdefault(current_year, {})
            current_sems = ["", ""]
            continue

        # ── Single-cell year row (fully merged) ──
        unique_vals = set(c.strip() for c in padded if c.strip())
        if len(unique_vals) == 1 and _YEAR_RE.search(c0):
            current_year = c0
            result.setdefault(current_year, {})
            current_sems = ["", ""]
            continue

        # ── Semester header row ──
        sem_left = bool(_SEM_RE.search(c0))
        sem_right = bool(_SEM_RE.search(c1))

        if sem_left or sem_right:
            if not current_year:
                current_year = "Year I"
                result.setdefault(current_year, {})

            if sem_left:
                current_sems[0] = c0
                result[current_year].setdefault(c0, [])
            if sem_right:
                current_sems[1] = c1
                result[current_year].setdefault(c1, [])
            continue

        # ── Data row: extract subjects ──
        if not current_year:
            continue

        for col_idx in range(min(len(padded), 2)):
            sem = current_sems[col_idx] if col_idx < len(current_sems) else ""
            if not sem:
                continue
            cell_text = padded[col_idx].strip()
            if not cell_text:
                continue
            subjects = _split_into_items(cell_text)
            result.setdefault(current_year, {}).setdefault(sem, []).extend(
                subjects
            )

    return result


def _parse_horizontal_table(rows: list[list[str]]) -> dict[str, str]:
    """Parse a 2-row label→value table into a flat dict.

    Header cells are normalised to ``lowercase_underscore`` keys.
    Merged / title rows are skipped before parsing.
    """
    filtered = _skip_merged_title_rows(rows)
    if len(filtered) < 2:
        return {}
    headers, values = filtered[0], filtered[1]
    out: dict[str, str] = {}
    for i, h in enumerate(headers):
        key = _normalize_header_key(h)
        val = values[i].strip() if i < len(values) else ""
        if key and val:
            out[key] = val
    return out


def _parse_uniform_table(rows: list[list[str]]) -> list[dict[str, str]]:
    """Parse a 2-column table with mostly-identical column-2 values.

    Still returned as a list of row objects (not collapsed).
    """
    if not rows:
        return []
    h0 = _normalize_header_key(rows[0][0]) or "col1"
    h1 = _normalize_header_key(rows[0][1]) or "col2"
    out: list[dict[str, str]] = []
    for row in rows[1:]:
        if len(row) < 2:
            continue
        v0, v1 = row[0].strip(), row[1].strip()
        if v0 or v1:
            out.append({h0: v0, h1: v1})
    return out


def _parse_standard_table(rows: list[list[str]]) -> list[dict[str, str]]:
    """Parse a standard table: first row = headers, rest = data dicts.

    Header keys are normalised to ``lowercase_underscore``.
    Empty rows are skipped entirely.
    """
    if not rows:
        return []

    headers = [_normalize_header_key(h) for h in rows[0]]
    # Fill unnamed columns with generic keys
    for i, h in enumerate(headers):
        if not h:
            headers[i] = f"col_{i + 1}"

    out: list[dict[str, str]] = []
    for row in rows[1:]:
        obj: dict[str, str] = {}
        has_content = False
        for i, h in enumerate(headers):
            val = row[i].strip() if i < len(row) else ""
            if val:
                has_content = True
            obj[h] = val
        if has_content:
            out.append(obj)
    return out


# ━━━━━━━━━━━━━━  STAGE 5 — ENVELOPE BUILDER  ━━━━━━━━━━━━━━


def _build_envelope(
    stripped_key: str,
    original_heading: str,
    content_type: str,
    table_subtype: str | None,
    content: Any,
) -> dict[str, Any]:
    """Wrap parsed content in the standard section envelope."""

    has_intro = False
    intro_text = None

    if content_type == "mixed" and isinstance(content, dict):
        raw_intro = content.get("intro_text")
        if raw_intro:
            has_intro = True
            intro_text = raw_intro

    return {
        "stripped_key": stripped_key,
        "original_heading": original_heading,
        "content_type": content_type,
        "table_subtype": table_subtype,
        "has_intro_text": has_intro,
        "intro_text": intro_text,
        "content": content,
    }


# ━━━━━━━━━━━━━━━━━━━━━━  SHARED HELPERS  ━━━━━━━━━━━━━━━━━━━━━━


def _extract_raw_text(doc) -> str:
    """Return the full document as a single plain-text string.

    Includes both paragraph text and table cell text so the downstream
    AI layer can scan for stats (duration, fee, EMI, etc.).
    """
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
    return "\n".join(parts)


def _texts_to_html(texts: list[str]) -> str:
    """Wrap plain-text items in ``<p>`` tags.

    Only uses ``<p>``, ``<ul>``, ``<li>``, ``<strong>``, ``<em>`` tags
    per the parser contract.
    """
    parts: list[str] = []
    for t in texts:
        escaped = (
            t.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        parts.append(f"<p>{escaped}</p>")
    return "\n".join(parts)


def _build_faq(texts: list[str]) -> list[dict[str, str]]:
    """Build FAQ pairs from alternating question / answer text items.

    Non-question items between pairs are absorbed into the preceding
    answer.
    """
    faq: list[dict[str, str]] = []
    i = 0
    while i < len(texts):
        q = texts[i].strip()
        if q.endswith("?"):
            a = texts[i + 1].strip() if i + 1 < len(texts) else ""
            faq.append({"question": q, "answer": a})
            i += 2
        else:
            # Absorb non-question text into previous answer
            if faq:
                faq[-1]["answer"] += " " + q
            i += 1
    return faq


def _split_into_items(text: str) -> list[str]:
    """Split a cell's text into individual items (by newline / bullet char).

    Used for extracting subject lists from semester-table data cells.
    """
    lines = text.split("\n")
    items: list[str] = []
    for line in lines:
        cleaned = _BULLET_RE.sub("", line).strip()
        if cleaned and len(cleaned) > 1:
            items.append(cleaned)
    return items if items else ([text.strip()] if text.strip() else [])


def _normalize_header_key(text: str) -> str:
    """Normalise a table header string into a ``lowercase_underscore`` key.

    - Removes ``**`` bold markers
    - Removes special characters except underscore
    - Spaces → underscores
    - Collapses multiple underscores
    - Strips leading/trailing underscores
    """
    t = text.strip()
    t = t.replace("**", "")
    t = _NON_ALNUM_RE.sub("", t)
    t = t.strip().replace(" ", "_").lower()
    t = _MULTI_USCORE_RE.sub("_", t)
    return t.strip("_")
